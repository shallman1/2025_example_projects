# analysis/fingerprint.py

import asyncio
import re
from collections import defaultdict
from docx import Document
from utils.api_utils import query_iris_api

def analyze_data(domains, limit_percentage=10, include_empty=False):
    total_domains = len(domains)
    attribute_counts = defaultdict(lambda: defaultdict(int))
    potential_dedicated_ips = set()

    for domain in domains:
        domain_name = domain.get('domain', '')
        analyze_domain(
            domain,
            attribute_counts,
            potential_dedicated_ips,
            domain_name,
            include_empty
        )

    high_correlation = get_high_correlation(
        attribute_counts,
        total_domains,
        limit_percentage
    )
    sorted_high_correlation = sort_attributes_by_highest_percentage(
        high_correlation
    )

    return sorted_high_correlation, potential_dedicated_ips

def normalize_attribute(attr):
    # Remove enumeration indices (numbers) from the attribute name
    attr = re.sub(r'_\d+', '', attr)
    return attr

def analyze_domain(
    domain,
    attribute_counts,
    potential_dedicated_ips,
    domain_name,
    include_empty
):
    excluded_fields = {'count', 'alexa', 'popularity_rank'}

    for attr, value in domain.items():
        if any(excluded in attr for excluded in excluded_fields):
            continue
        # Process only keys that end with '_value'
        if attr.endswith('_value'):
            # Get the base attribute name
            base_attr = attr[:-6]  # Remove '_value'
            # Normalize attribute to remove enumeration indices
            base_attr = normalize_attribute(base_attr)
            count_key = base_attr + '_count'
            count_value = domain.get(count_key, '')
            # Convert value to string for comparison
            value_str = str(value) if value is not None else ''
            # Skip empty values unless include_empty is True
            if not include_empty and (not value_str or value_str.lower() == 'empty'):
                continue
            # Replace underscores with spaces in context key for readability
            context_key = f"{base_attr.replace('_', ' ')}: {value_str}"
            # Store the count value
            try:
                count_value = int(count_value) if count_value else 0
            except ValueError:
                count_value = 0

            # Update attribute counts
            attribute_counts[base_attr][context_key] += 1

            # Collect potential dedicated IPs
            if base_attr == 'ip address':
                if 0 <= count_value <= 20:
                    potential_dedicated_ips.add(value_str)

def get_high_correlation(attribute_counts, total_domains, limit_percentage):
    high_correlation = {}
    for attr, values in attribute_counts.items():
        correlations = []
        for context_key, count in values.items():
            percentage = min((count / total_domains) * 100, 100)
            if percentage >= limit_percentage:
                correlations.append((context_key, percentage))
        if correlations:
            high_correlation[attr] = correlations
    return high_correlation

def sort_attributes_by_highest_percentage(data):
    sorted_data = {}
    attributes_with_max_percentage = []
    for attr, values in data.items():
        max_percentage = max(values, key=lambda x: x[1])[1]
        attributes_with_max_percentage.append((attr, max_percentage))
    attributes_with_max_percentage.sort(key=lambda x: x[1], reverse=True)
    for attr, _ in attributes_with_max_percentage:
        sorted_data[attr] = sorted(data[attr], key=lambda x: x[1], reverse=True)
    return sorted_data

def generate_report(high_correlation, potential_dedicated_ips):
    document = Document()

    # Add title
    document.add_heading('Correlation Report', 0)

    # High Correlation Attributes Section
    document.add_heading('1. High Correlation Attributes', level=1)
    for attr, values in high_correlation.items():
        document.add_heading(attr.replace('_', ' '), level=2)

        for context_key, percentage in values:
            if not context_key.strip(", "):
                continue
            p = document.add_paragraph(style='List Bullet')
            p.add_run(f'{context_key}: {percentage:.2f}%')

    # Potential Dedicated IPs Section
    document.add_heading("2. Potential Dedicated IP's", level=1)
    if potential_dedicated_ips:
        for ip in sorted(potential_dedicated_ips):
            p = document.add_paragraph(style='List Bullet')
            p.add_run(ip)
    else:
        p = document.add_paragraph("No potential dedicated IPs found.")

    return document

def run_analysis(search_hash, limit_percentage, iris_key, iris_user, include_empty=False):
    loop = asyncio.new_event_loop()
    asyncio.set_event_loop(loop)
    domains = loop.run_until_complete(
        query_iris_api(iris_key, iris_user, search_hash)
    )
    high_correlation, potential_dedicated_ips = analyze_data(
        domains,
        limit_percentage,
        include_empty
    )
    document = generate_report(
        high_correlation,
        potential_dedicated_ips
    )
    return document



